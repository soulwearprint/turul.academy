"""
First parse pass over the official NAT 2020 Fizika kerettanterv docx files.

Unlike History (docx encodes Témák as table rows, elements curated separately),
Physics's docx has NO Téma-level breakdown at all — each Témakör is a flat block of:
  Tanulási eredmények        (learning outcomes, prose)
  Fejlesztési feladatok és ismeretek  (mandatory facts, bullet list)
  Fogalmak                   (concepts, comma/semicolon-separated list)
  Javasolt tevékenységek      (suggested activities/experiments, bullet list)
This pass extracts that structure directly and faithfully — no clustering, no LLM,
no invented Téma titles. Téma derivation happens in a later LLM pass
(cluster_temak_physics.py) that reads this file's output.

Scope: grades 7-10 only (confirmed with user 2026-07-08) — Fizika_F.docx covers
7-8 (10 Témakör), Fizika_K.docx covers 9-10 (15 Témakör). No grade 11-12 source
document exists; that gap is out of scope for this re-foundation.

Output: content/nat_curriculum/physics_nat2020.json
  { "source": "...",
    "altalanos_iskola": {"7-8": [ {temakor, oraszam, tanulasi_eredmenyek:[...],
                                   fejlesztesi_feladatok:[...], fogalmak:[...],
                                   tevekenysegek:[...]} , ... ]},
    "gimnazium":        {"9-10": [ ... ]} }
"""
import os, re, json
import docx

HERE = os.path.dirname(__file__)
DOCX = {
    "altalanos_iskola": {"path": "/Users/gabor/Documents/Turul/Turul-Academy/Physics/Fizika_F.docx", "band": "7-8"},
    "gimnazium":        {"path": "/Users/gabor/Documents/Turul/Turul-Academy/Physics/Fizika_K.docx", "band": "9-10"},
}
OUT = os.path.join(HERE, "../nat_curriculum/physics_nat2020.json")

H3_TANULASI = "tanulási eredmények"
H3_FELADATOK = "fejlesztési feladatok és ismeretek"
H3_FOGALMAK = "fogalmak"
H3_TEVEKENYSEGEK = "javasolt tevékenységek"


def _oraszam_list(path):
    """Overview table hours, in row order. Skips the summary row.
    Matched to Témakör sections POSITIONALLY, not by title — titles occasionally
    drift slightly between the overview table and the section header in the source
    (e.g. F's table says "Víz, levegő és szilárd anyagok a háztartásban..." but the
    section header says "Víz és levegő a háztartásban...") while row/section order
    is reliably 1:1."""
    d = docx.Document(path)
    out = []
    for row in d.tables[0].rows[1:]:
        title = row.cells[0].text.strip()
        hours = row.cells[1].text.strip()
        if not title or not hours.isdigit():
            continue
        out.append(int(hours))
    return out


def _parse_docx(path, band):
    hours_list = _oraszam_list(path)
    d = docx.Document(path)

    topics = []
    cur = None
    section = None  # which H3 bucket we're currently collecting into

    def flush():
        if cur is not None:
            topics.append(cur)

    for p in d.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        style = (p.style.name or "").lower()

        if t.startswith("Témakör:"):
            flush()
            title = t.split("Témakör:", 1)[1].strip()
            idx = len(topics)
            cur = {"temakor": title, "oraszam": hours_list[idx] if idx < len(hours_list) else 0,
                   "tanulasi_eredmenyek": [], "fejlesztesi_feladatok": [],
                   "fogalmak": [], "tevekenysegek": []}
            section = None
            continue
        if cur is None:
            continue  # front-matter before the first Témakör

        if t.startswith("Javasolt óraszám"):
            continue  # already captured from the overview table

        # The Fogalmak list is exactly one paragraph immediately after the "Fogalmak"
        # heading, inconsistently styled "Heading 3" or "Normal" across topics — so
        # check this before the heading-name switch, regardless of style.
        if section == "fogalmak" and not cur["fogalmak"]:
            cur["fogalmak"] = [x.strip() for x in re.split(r"[,;]", t) if x.strip()]
            continue

        if style == "heading 3":
            low = t.strip().lower()
            if low == H3_TANULASI:
                section = "tanulasi_eredmenyek"
            elif low == H3_FELADATOK:
                section = "fejlesztesi_feladatok"
            elif low == H3_FOGALMAK:
                section = "fogalmak"
            elif low == H3_TEVEKENYSEGEK:
                section = "tevekenysegek"
            continue

        if section in ("tanulasi_eredmenyek", "fejlesztesi_feladatok", "tevekenysegek"):
            cur[section].append(t)
        # Normal-style prose before "A témakör tanulása eredményeként..." etc. inside
        # tanulasi_eredmenyek is intentionally kept — it's part of the outcomes text.

    flush()
    return topics


def main():
    result = {"source": "NAT 2020 (2024 kerettanterv), Fizika — user-provided Fizika_F/K.docx, grades 7-10 only"}
    for section, info in DOCX.items():
        path, band = info["path"], info["band"]
        if not os.path.exists(path):
            print(f"⚠ missing {path}"); continue
        topics = _parse_docx(path, band)
        result[section] = {band: topics}
        print(f"{section} [{band}]: {len(topics)} Témakör")
        for tp in topics:
            print(f"  - {tp['temakor']}  ({tp['oraszam']}h)  "
                  f"feladatok={len(tp['fejlesztesi_feladatok'])} fogalmak={len(tp['fogalmak'])} "
                  f"tevekenysegek={len(tp['tevekenysegek'])}")

    json.dump(result, open(OUT, "w", encoding="utf-8"), ensure_ascii=False, indent=2)
    print(f"\n→ {OUT}")


if __name__ == "__main__":
    main()
